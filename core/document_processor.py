"""
DataNeuron Multi-Format Document Processing Module
==================================================

This module provides document processing capabilities for various file formats (PDF, TXT, DOCX).
It extracts text content and metadata from documents in a standardized format.

Features:
- Supports PDF, TXT, and DOCX file formats
- Extracts comprehensive metadata (file info, creation date, etc.)
- File size and extension validation
- Robust error handling with detailed logging
- Strategy pattern for extensible format support

Usage:
    from core.document_processor import DocumentProcessor
    
    processor = DocumentProcessor()
    document = processor.process("/path/to/file.pdf")
    
    if document:
        print(f"Content: {document.content}")
        print(f"Metadata: {document.metadata}")
"""

import os
import sys
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any

# Add project root to Python path for imports
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from utils.logger import logger
from config.settings import SUPPORTED_EXTENSIONS, MAX_FILE_SIZE_MB

# Document format specific imports
try:
    import PyPDF2
except ImportError:
    logger.warning("PyPDF2 not installed. PDF processing will not be available.")
    PyPDF2 = None

try:
    import docx
except ImportError:
    logger.warning("python-docx not installed. DOCX processing will not be available.")
    docx = None


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class Document:
    """
    Represents a processed document with content and metadata.
    
    Attributes:
        content: Extracted text content from the document
        metadata: Dictionary containing file information and format-specific data
    """
    content: str
    metadata: Dict[str, Any]


# ============================================================================
# ABSTRACT BASE CLASS FOR DOCUMENT READERS
# ============================================================================

class DocumentReader(ABC):
    """
    Abstract base class for document readers.
    
    Each file format should have its own reader class that inherits from this.
    """
    
    @abstractmethod
    def read(self, file_path: Path) -> Optional[Document]:
        """
        Read and process a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with content and metadata, or None if processing failed
        """
        pass


# ============================================================================
# CONCRETE READER IMPLEMENTATIONS
# ============================================================================

class TxtReader(DocumentReader):
    """Reader for plain text files (.txt)."""
    
    def read(self, file_path: Path) -> Optional[Document]:
        """
        Read a text file and extract its content.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Document object or None if reading failed
        """
        try:
            logger.info(f"Reading TXT file: {file_path}")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                logger.error(f"Failed to decode TXT file with any encoding: {file_path}")
                return None
            
            # Get file metadata
            file_stat = file_path.stat()
            metadata = self._get_base_metadata(file_path, file_stat)
            metadata.update({
                'encoding': used_encoding,
                'line_count': content.count('\n') + 1,
                'character_count': len(content),
                'word_count': len(content.split())
            })
            
            logger.success(f"Successfully read TXT file: {len(content)} characters")
            return Document(content=content, metadata=metadata)
            
        except FileNotFoundError:
            logger.error(f"TXT file not found: {file_path}")
            return None
        except Exception as e:
            logger.exception(f"Error reading TXT file {file_path}: {str(e)}")
            return None
    
    def _get_base_metadata(self, file_path: Path, file_stat) -> Dict[str, Any]:
        """Extract base metadata common to all file types."""
        return {
            'file_name': file_path.name,
            'file_path': str(file_path.absolute()),
            'file_type': file_path.suffix.lower(),
            'file_size_mb': round(file_stat.st_size / (1024 * 1024), 3),
            'creation_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modification_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        }


class PdfReader(DocumentReader):
    """Reader for PDF files (.pdf)."""
    
    def read(self, file_path: Path) -> Optional[Document]:
        """
        Read a PDF file and extract its text content.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Document object or None if reading failed
        """
        if PyPDF2 is None:
            logger.error("PyPDF2 not available. Cannot process PDF files.")
            return None
        
        try:
            logger.info(f"Reading PDF file: {file_path}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.error(f"PDF file is encrypted: {file_path}")
                    return None
                
                # Extract text from all pages
                content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                
                # Get file metadata
                file_stat = file_path.stat()
                metadata = self._get_base_metadata(file_path, file_stat)
                
                # Add PDF-specific metadata
                try:
                    pdf_info = pdf_reader.metadata
                    metadata.update({
                        'page_count': len(pdf_reader.pages),
                        'character_count': len(content),
                        'word_count': len(content.split()),
                        'pdf_title': str(pdf_info.get('/Title', '')) if pdf_info else '',
                        'pdf_author': str(pdf_info.get('/Author', '')) if pdf_info else '',
                        'pdf_subject': str(pdf_info.get('/Subject', '')) if pdf_info else '',
                        'pdf_creator': str(pdf_info.get('/Creator', '')) if pdf_info else '',
                    })
                except Exception as e:
                    logger.warning(f"Error extracting PDF metadata: {str(e)}")
                    metadata.update({
                        'page_count': len(pdf_reader.pages),
                        'character_count': len(content),
                        'word_count': len(content.split()),
                    })
                
                logger.success(f"Successfully read PDF: {metadata['page_count']} pages, {len(content)} characters")
                return Document(content=content, metadata=metadata)
                
        except FileNotFoundError:
            logger.error(f"PDF file not found: {file_path}")
            return None
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error for file {file_path}: {str(e)}")
            return None
        except Exception as e:
            logger.exception(f"Error reading PDF file {file_path}: {str(e)}")
            return None
    
    def _get_base_metadata(self, file_path: Path, file_stat) -> Dict[str, Any]:
        """Extract base metadata common to all file types."""
        return {
            'file_name': file_path.name,
            'file_path': str(file_path.absolute()),
            'file_type': file_path.suffix.lower(),
            'file_size_mb': round(file_stat.st_size / (1024 * 1024), 3),
            'creation_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modification_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        }


class DocxReader(DocumentReader):
    """Reader for Microsoft Word documents (.docx)."""
    
    def read(self, file_path: Path) -> Optional[Document]:
        """
        Read a DOCX file and extract its text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Document object or None if reading failed
        """
        if docx is None:
            logger.error("python-docx not available. Cannot process DOCX files.")
            return None
        
        try:
            logger.info(f"Reading DOCX file: {file_path}")
            
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Get file metadata
            file_stat = file_path.stat()
            metadata = self._get_base_metadata(file_path, file_stat)
            
            # Add DOCX-specific metadata
            try:
                core_props = doc.core_properties
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'character_count': len(content),
                    'word_count': len(content.split()),
                    'docx_title': core_props.title or '',
                    'docx_author': core_props.author or '',
                    'docx_subject': core_props.subject or '',
                    'docx_created': core_props.created.isoformat() if core_props.created else '',
                    'docx_modified': core_props.modified.isoformat() if core_props.modified else '',
                })
            except Exception as e:
                logger.warning(f"Error extracting DOCX metadata: {str(e)}")
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'character_count': len(content),
                    'word_count': len(content.split()),
                })
            
            logger.success(f"Successfully read DOCX: {metadata['paragraph_count']} paragraphs, {len(content)} characters")
            return Document(content=content, metadata=metadata)
            
        except FileNotFoundError:
            logger.error(f"DOCX file not found: {file_path}")
            return None
        except docx.opc.exceptions.PackageNotFoundError:
            logger.error(f"Invalid DOCX file format: {file_path}")
            return None
        except Exception as e:
            logger.exception(f"Error reading DOCX file {file_path}: {str(e)}")
            return None
    
    def _get_base_metadata(self, file_path: Path, file_stat) -> Dict[str, Any]:
        """Extract base metadata common to all file types."""
        return {
            'file_name': file_path.name,
            'file_path': str(file_path.absolute()),
            'file_type': file_path.suffix.lower(),
            'file_size_mb': round(file_stat.st_size / (1024 * 1024), 3),
            'creation_date': datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modification_date': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
        }


# ============================================================================
# MAIN DOCUMENT PROCESSOR CLASS
# ============================================================================

class DocumentProcessor:
    """
    Main document processor that handles multiple file formats.
    
    This class validates files and delegates processing to appropriate readers.
    """
    
    def __init__(self):
        """Initialize the document processor with configured settings."""
        self.supported_extensions = SUPPORTED_EXTENSIONS
        self.max_file_size_mb = MAX_FILE_SIZE_MB
        
        # Initialize readers for each supported format
        self.readers = {
            '.txt': TxtReader(),
            '.pdf': PdfReader(),
            '.docx': DocxReader(),
        }
        
        logger.info(f"DocumentProcessor initialized with extensions: {self.supported_extensions}")
        logger.info(f"Maximum file size: {self.max_file_size_mb} MB")
    
    def process(self, file_path: str) -> Optional[Document]:
        """
        Process a document file and extract its content and metadata.
        
        Args:
            file_path: Path to the document file (string)
            
        Returns:
            Document object with content and metadata, or None if processing failed
        """
        try:
            path = Path(file_path)
            
            # Validate file existence
            if not path.exists():
                logger.error(f"File does not exist: {file_path}")
                return None
            
            if not path.is_file():
                logger.error(f"Path is not a file: {file_path}")
                return None
            
            # Validate file extension
            file_extension = path.suffix.lower()
            if file_extension not in self.supported_extensions:
                logger.warning(f"Unsupported file extension '{file_extension}': {file_path}")
                logger.warning(f"Supported extensions: {self.supported_extensions}")
                return None
            
            # Validate file size
            file_size_mb = path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                logger.error(f"File size ({file_size_mb:.2f} MB) exceeds limit ({self.max_file_size_mb} MB): {file_path}")
                return None
            
            # Get appropriate reader and process the file
            reader = self.readers.get(file_extension)
            if not reader:
                logger.error(f"No reader available for extension '{file_extension}': {file_path}")
                return None
            
            logger.info(f"Processing file: {file_path} ({file_size_mb:.2f} MB)")
            document = reader.read(path)
            
            if document:
                logger.success(f"Successfully processed document: {path.name}")
                return document
            else:
                logger.error(f"Failed to process document: {path.name}")
                return None
                
        except Exception as e:
            logger.exception(f"Unexpected error processing file {file_path}: {str(e)}")
            return None


# ============================================================================
# TESTING AND VALIDATION
# ============================================================================

if __name__ == "__main__":
    """
    Test the document processor with various scenarios.
    This helps verify that all functionality is working correctly.
    """
    
    print("=== DataNeuron Document Processor Test ===")
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Test scenarios
    test_cases = [
        {
            'name': 'Non-existent file',
            'path': 'non_existent_file.pdf',
            'expected': None
        },
        {
            'name': 'Unsupported extension',
            'path': __file__.replace('.py', '.xyz'),  # Create a fake extension
            'expected': None
        },
        {
            'name': 'Current Python file (should work as .py is treated as text)',
            'path': __file__,
            'expected': None  # .py is not in supported extensions
        }
    ]
    
    # Create test files for demonstration
    test_dir = Path("test_documents")
    test_dir.mkdir(exist_ok=True)
    
    # Create a simple text file
    txt_file = test_dir / "test.txt"
    with open(txt_file, 'w', encoding='utf-8') as f:
        f.write("This is a test document.\nIt contains multiple lines.\nFor testing purposes.")
    
    test_cases.append({
        'name': 'Valid text file',
        'path': str(txt_file),
        'expected': 'Document'
    })
    
    # Create a large dummy file to test size limit
    large_file = test_dir / "large.txt"
    large_content = "x" * (MAX_FILE_SIZE_MB * 1024 * 1024 + 1000)  # Exceed size limit
    with open(large_file, 'w') as f:
        f.write(large_content)
    
    test_cases.append({
        'name': 'File exceeding size limit',
        'path': str(large_file),
        'expected': None
    })
    
    # Run tests
    print(f"\nRunning {len(test_cases)} test cases:\n")
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"Test {i}: {test_case['name']}")
        print(f"  File: {test_case['path']}")
        
        result = processor.process(test_case['path'])
        
        if test_case['expected'] is None:
            if result is None:
                print("   PASS - Correctly returned None")
            else:
                print("  L FAIL - Expected None but got Document")
        else:
            if result is not None:
                print("   PASS - Successfully processed document")
                print(f"     Content length: {len(result.content)} characters")
                print(f"     Metadata keys: {list(result.metadata.keys())}")
            else:
                print("  L FAIL - Expected Document but got None")
        
        print()
    
    # Cleanup test files
    try:
        txt_file.unlink(missing_ok=True)
        large_file.unlink(missing_ok=True)
        test_dir.rmdir()
        logger.info("Test files cleaned up successfully")
    except Exception as e:
        logger.warning(f"Could not clean up test files: {e}")
    
    print("=== Test Complete ===")
    print("Check the logs above for detailed processing information.")